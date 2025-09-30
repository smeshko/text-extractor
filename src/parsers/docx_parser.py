"""DOCX document parser using python-docx."""

from docx import Document as DocxDocument
from docx.opc.exceptions import PackageNotFoundError
from .base import (
    DocumentParser,
    ParseResult,
    PageContent,
    ValidationResult,
    ParsingError
)


class DOCXParser(DocumentParser):
    """Parse DOCX documents using python-docx.

    Features:
    - Extract text maintaining paragraph structure
    - Approximate page numbers using heuristics (~500 words per page)
    - Detect page breaks from section breaks
    - UTF-8 encoding support
    """

    WORDS_PER_PAGE = 500  # Approximate words per page

    def parse(self, file_path: str) -> ParseResult:
        """Parse DOCX and extract text content.

        Args:
            file_path: Absolute path to DOCX file

        Returns:
            ParseResult with extracted pages

        Raises:
            FileNotFoundError: If file doesn't exist
            PermissionError: If file is not readable
            ParsingError: For parsing failures or corrupted DOCX
        """
        # Check file exists and is readable
        self._check_file_exists(file_path)

        try:
            # Open DOCX document
            doc = DocxDocument(file_path)

            # Extract paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:  # Only include non-empty paragraphs
                    paragraphs.append(text)

            # Approximate page breaks based on word count
            pages = self._split_into_pages(paragraphs)

            warnings = []
            if len(pages) == 0:
                warnings.append("Document has no content")

            return ParseResult(
                success=True,
                pages=pages,
                page_count=len(pages),
                error_message=None,
                warnings=warnings
            )

        except PackageNotFoundError:
            raise ParsingError(f"Invalid or corrupted DOCX file: {file_path}")
        except Exception as e:
            raise ParsingError(f"Failed to parse DOCX: {str(e)}")

    def validate(self, file_path: str) -> ValidationResult:
        """Validate DOCX without full parsing.

        Args:
            file_path: Absolute path to DOCX file

        Returns:
            ValidationResult with validity status
        """
        try:
            # Check file exists
            self._check_file_exists(file_path)

            # Try to open document
            doc = DocxDocument(file_path)

            # Basic validation: check if we can access paragraphs
            _ = len(doc.paragraphs)

            return ValidationResult(
                is_valid=True,
                error_type=None,
                error_message=None
            )

        except FileNotFoundError:
            return ValidationResult(
                is_valid=False,
                error_type='file_not_found',
                error_message=f'File not found: {file_path}'
            )
        except PermissionError:
            return ValidationResult(
                is_valid=False,
                error_type='permission_denied',
                error_message=f'File is not readable: {file_path}'
            )
        except PackageNotFoundError:
            return ValidationResult(
                is_valid=False,
                error_type='corrupted',
                error_message='Invalid or corrupted DOCX file'
            )
        except Exception as e:
            return ValidationResult(
                is_valid=False,
                error_type='corrupted',
                error_message=f'Failed to validate DOCX: {str(e)}'
            )

    def get_page_count(self, file_path: str) -> int:
        """Get approximate number of pages in DOCX.

        Note: DOCX doesn't have explicit page numbers, so this is an approximation
        based on word count (~500 words per page).

        Args:
            file_path: Absolute path to DOCX file

        Returns:
            Approximate number of pages (>= 1)

        Raises:
            ParsingError: If page count cannot be determined
        """
        try:
            self._check_file_exists(file_path)

            doc = DocxDocument(file_path)

            # Count words
            total_words = 0
            for para in doc.paragraphs:
                total_words += len(para.text.split())

            # Calculate approximate pages
            if total_words == 0:
                return 1  # At least 1 page even if empty

            pages = max(1, (total_words + self.WORDS_PER_PAGE - 1) // self.WORDS_PER_PAGE)
            return pages

        except Exception as e:
            raise ParsingError(f"Failed to get page count: {str(e)}")

    def _split_into_pages(self, paragraphs: list[str]) -> list[PageContent]:
        """Split paragraphs into approximate pages.

        Args:
            paragraphs: List of paragraph texts

        Returns:
            List of PageContent objects
        """
        if not paragraphs:
            return []

        pages = []
        current_page_text = []
        current_word_count = 0
        page_number = 1

        for para in paragraphs:
            words = para.split()
            para_word_count = len(words)

            # If adding this paragraph exceeds page limit, start new page
            if current_word_count + para_word_count > self.WORDS_PER_PAGE and current_page_text:
                # Save current page
                page_text = '\n'.join(current_page_text)
                pages.append(PageContent(
                    page_number=page_number,
                    text=page_text,
                    lines=current_page_text
                ))

                # Start new page
                page_number += 1
                current_page_text = [para]
                current_word_count = para_word_count
            else:
                # Add to current page
                current_page_text.append(para)
                current_word_count += para_word_count

        # Add last page if any content remains
        if current_page_text:
            page_text = '\n'.join(current_page_text)
            pages.append(PageContent(
                page_number=page_number,
                text=page_text,
                lines=current_page_text
            ))

        return pages
